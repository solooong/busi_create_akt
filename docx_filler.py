import os
import logging
from typing import Dict, List, Optional
from docx import Document
import re  # <--- ДОБАВИТЬ ЭТУ СТРОКУ
import copy  # <--- ДОБАВИТЬ ЭТУ СТРОКУ

logger = logging.getLogger(__name__)

# Поля, относящиеся к договору в целом, а не к конкретному участнику
CONTRACT_KEYS = {
    'договор_№', 'дата_договора', 'название_договора',
    'кадастровый_номер', 'адрес_строительства',
    'номер_квартиры', 'этаж', 'общая_площадь',
    'жилая_площадь', 'площадь_балкона', 'тип_квартиры'
}


class DocxFiller:
    """Заполняет шаблон акта приёма-передачи данными участников."""

    def __init__(self, template_path: str):
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")
        self.template_path = template_path

    # ------------------------------------------------------------------
    # Публичные методы
    # ------------------------------------------------------------------
    def fill_single_act(self, participant: Dict, output_path: str) -> str:
        """Заполняет акт для одного участника (без группировки)."""
        doc = Document(self.template_path)
        flat = self._flatten_participant(participant)
        flat['количество_сторон'] = '2'
        # Явно гарантируем наличие ФИО для замены в преамбуле
        flat['ФИО'] = flat.get('ФИО', '-')
        self._replace_all(doc, flat)
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        doc.save(output_path)
        logger.info(f"Акт сохранён: {output_path}")
        return output_path

    def fill_multiple_acts(self, participants: List[Dict], output_dir: str,
                           prefix: str = "Акт_") -> List[str]:
        """Индивидуальные акты для каждого участника."""
        created = []
        for i, p in enumerate(participants, 1):
            num = p.get('доп_данные', {}).get('договор_№', f'уч{i}')
            name = p.get('ФИО', f'участник_{i}')
            safe_name = "".join(c for c in name if c.isalnum() or c in ' _-')
            filename = f"{prefix}{num}_{safe_name}.docx"
            path = os.path.join(output_dir, filename)
            try:
                created.append(self.fill_single_act(p, path))
            except Exception as e:
                logger.error(f"Не удалось создать акт для {name}: {e}")
        return created

    def fill_grouped_acts(self, participants: List[Dict], output_dir: str,
                          group_by: str = 'договор_№',
                          prefix: str = 'Акт_договор_') -> List[str]:
        """
        Группирует участников по номеру договора и создаёт по одному акту на договор.
        """
        groups = {}
        for p in participants:
            key = p.get('доп_данные', {}).get(group_by, 'без_договора')
            groups.setdefault(key, []).append(p)

        created = []
        for contract_num, group in groups.items():
            safe_num = "".join(c for c in str(contract_num) if c.isalnum() or c in ' _-')
            filename = f"{prefix}{safe_num}.docx"
            path = os.path.join(output_dir, filename)
            try:
                self._fill_group_act(group, path)
                created.append(path)
            except Exception as e:
                logger.error(f"Ошибка акта для договора {contract_num}: {e}")
        return created

    # ------------------------------------------------------------------
    # Внутренние методы
    # ------------------------------------------------------------------
    def _fill_group_act(self, group: List[Dict], output_path: str):
        doc = Document(self.template_path)

        # 1. Сначала заполняем таблицу с деталями участников, пока плейсхолдеры на месте
        self._fill_owners_table_recursive(doc, group)

        # 2. Затем заменяем все остальные плейсхолдеры, включая {{ ФИО }} в преамбуле
        common = self._get_contract_common(group[0])
        common['количество_сторон'] = str(len(group) + 1)
        
        # Добавляем ФИО как строку для преамбулы (список всех участников через запятую)
        common['ФИО'] = ", ".join([p.get('ФИО', '-') for p in group])
        
        self._replace_all(doc, common)

        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        doc.save(output_path)
        logger.info(f"Групповой акт сохранён: {output_path}")

    def _flatten_participant(self, p: Dict) -> Dict:
        """Собирает все данные участника в плоский словарь."""
        flat = {
            'ФИО': p.get('ФИО') or '-',
            'телефон': p.get('телефон') or '-',
            'email': p.get('email') or '-',
        }
        for k, v in p.get('доп_данные', {}).items():
            flat[k] = v if v and str(v).strip() else '-'
        return flat

    def _get_contract_common(self, p: Dict) -> Dict:
        """Возвращает только общие для договора поля."""
        dop = p.get('доп_данные', {})
        common = {}
        for key in CONTRACT_KEYS:
            val = dop.get(key, '')
            common[key] = val if val and str(val).strip() else '-'
        return common

    # ------------------------------------------------------------------
    # Рекурсивный поиск таблицы с реквизитами и её заполнение
    # ------------------------------------------------------------------
    def _fill_owners_table_recursive(self, doc: Document, group: List[Dict]):
        """Ищет таблицу реквизитов участников во всех таблицах документа и заполняет её."""
        for table in doc.tables:
            if self._try_fill_table(table, group):
                return True
        
        # Если не нашли в верхнем уровне, ищем во вложенных таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for nested_table in cell.tables:
                        if self._try_fill_table(nested_table, group):
                            return True
        return False

    def _try_fill_table(self, table, group: List[Dict]) -> bool:
        """Пытается найти образцовую строку в конкретной таблице и заполнить её."""
        
        # 1. Строгая проверка: это должна быть таблица реквизитов.
        detail_keys = ['дата_рождения', 'паспорт', 'инн', 'снилс', 'адрес_проживания', 'телефон', 'email']
        has_detail_placeholder = False
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.replace(' ', '').lower()
                if any(f'{{{{{key}}}}}' in text or f'{{{key}}}' in text for key in detail_keys):
                    has_detail_placeholder = True
                    break
            if has_detail_placeholder:
                break
        
        if not has_detail_placeholder:
            return False

        # 2. Поиск образцовой строки
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.replace(' ', '').lower()
                if '{{фио' in text or '{{дата_рождения' in text or '{{паспорт' in text:
                    template_row = row
                    
                    # 3. ГЛАВНОЕ ИСПРАВЛЕНИЕ: Клонируем XML строки для остальных участников 
                    # ДО замены плейсхолдеров! Это сохранит все абзацы и форматирование.
                    cloned_rows_xml = []
                    for _ in group[1:]:
                        new_tr = copy.deepcopy(template_row._tr)
                        cloned_rows_xml.append(new_tr)
                    
                    # Вставляем копии после template_row в правильном порядке
                    for new_tr in reversed(cloned_rows_xml):
                        template_row._tr.addnext(new_tr)
                    
                    # Находим индекс template_row в обновленном списке строк таблицы
                    template_idx = -1
                    for i, r in enumerate(table.rows):
                        if r._tr is template_row._tr:
                            template_idx = i
                            break
                    
                    # 4. Заполняем строки данными участников
                    # template_row -> group[0]
                    self._fill_row(template_row, group[0])
                    
                    # cloned rows -> group[1:]
                    for i, p in enumerate(group[1:]):
                        row_idx = template_idx + 1 + i
                        if row_idx < len(table.rows):
                            self._fill_row(table.rows[row_idx], p)
                            
                    return True
        return False
    
    
    def _fill_row(self, row, participant: Dict):
        """Заменяет плейсхолдеры в строке таблицы данными одного участника."""
        flat = self._flatten_participant(participant)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                self._replace_in_paragraph(paragraph, flat)

    # ------------------------------------------------------------------
    # Универсальная замена плейсхолдеров
    # ------------------------------------------------------------------
    def _replace_all(self, doc: Document, replacements: Dict):
        """Заменяет плейсхолдеры во всех параграфах и таблицах документа."""
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
                    # Рекурсивно обрабатываем вложенные таблицы
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for paragraph in nested_cell.paragraphs:
                                    self._replace_in_paragraph(paragraph, replacements)

    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        """Заменяет все {{ключ}} в параграфе, сохраняя форматирование."""
        full_text = paragraph.text
        if '{{' not in full_text:
            return

        for key, value in replacements.items():
            placeholder = f'{{{{{key}}}}}'
            safe_value = str(value).strip() if value and str(value).strip() else "-"
            full_text = full_text.replace(placeholder, safe_value)

        if full_text != paragraph.text:
            if paragraph.runs:
                first_run = paragraph.runs[0]
                bold = first_run.font.bold
                italic = first_run.font.italic
                size = first_run.font.size
            else:
                bold = italic = size = None

            paragraph.clear()
            run = paragraph.add_run(full_text)
            if bold is not None:
                run.font.bold = bold
            if italic is not None:
                run.font.italic = italic
            if size is not None:
                run.font.size = size